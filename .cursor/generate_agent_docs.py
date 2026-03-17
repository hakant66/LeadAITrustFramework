#!/usr/bin/env python3
"""
Generate Agent Capabilities Documentation as .docx

This script creates a Word document (.docx) with comprehensive documentation
of all agents in the multi-agent system.
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is not installed.")
    print("Install it with: pip install python-docx")
    exit(1)

from datetime import datetime


def add_heading_with_style(doc, text, level=1, color=None):
    """Add a heading with optional color"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_table_from_dict(doc, data, headers=None):
    """Add a table from dictionary data"""
    if headers is None:
        headers = list(data[0].keys()) if data else []
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            if isinstance(row_data, dict):
                row_cells[i].text = str(row_data.get(header, ''))
            else:
                row_cells[i].text = str(row_data[i] if i < len(row_data) else '')
    
    return table


def create_agent_documentation():
    """Create the agent capabilities documentation"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Multi-Agent System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Agent Capabilities Documentation', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('LeadAI Trust Framework')
    doc.add_page_break()
    
    # Table of Contents placeholder
    toc_heading = doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Overview',
        '2. Agent Roster',
        '   2.1 The Coder',
        '   2.2 The Cleaner',
        '   2.3 The Verifier',
        '   2.4 The Detective',
        '   2.5 The EU Compliance Agent',
        '   2.6 The Coordinator',
        '3. Model Comparison',
        '4. Recommended Workflows',
        '5. Usage Guidelines'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Overview
    add_heading_with_style(doc, '1. Overview', 1)
    doc.add_paragraph(
        'This document provides comprehensive documentation of all agents in the '
        'LeadAI Trust Framework multi-agent development system. Each agent has '
        'specialized capabilities and uses different AI models optimized for '
        'specific tasks.'
    )
    
    # 2. Agent Roster
    add_heading_with_style(doc, '2. Agent Roster', 1)
    
    # 2.1 The Coder
    add_heading_with_style(doc, '2.1 The Coder (Agent A - Llama 3.1:8b)', 2)
    doc.add_paragraph('Model: Local Llama 3.1:8b via Ollama', style='Intense Quote')
    doc.add_paragraph('Speed: Fast (local, GPU-accelerated)', style='Intense Quote')
    
    doc.add_paragraph('Core Capabilities:', style='Heading 3')
    capabilities = [
        'Rapid code generation and implementation',
        'Writes new features and functionality',
        'Implements API endpoints, services, and components',
        'Follows project architecture patterns',
        'Delivers working code quickly'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('Best For:', style='Heading 3')
    best_for = [
        'New feature development',
        'API endpoint creation',
        'Quick prototypes',
        'Fast iterations'
    ]
    for item in best_for:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Guidelines:', style='Heading 3')
    guidelines = [
        'Speed over perfection',
        'Functionality first',
        'Follow existing patterns',
        'Quick incremental progress'
    ]
    for guideline in guidelines:
        doc.add_paragraph(guideline, style='List Bullet')
    
    doc.add_paragraph('Trigger: "Act as The Coder" or "I need code for..."', style='Intense Quote')
    
    # 2.2 The Cleaner
    add_heading_with_style(doc, '2.2 The Cleaner (Agent B - Claude 3.5 Sonnet)', 2)
    doc.add_paragraph('Model: Claude 3.5 Sonnet (via Cursor)', style='Intense Quote')
    doc.add_paragraph('Focus: Code quality and optimization', style='Intense Quote')
    
    doc.add_paragraph('Core Capabilities:', style='Heading 3')
    capabilities = [
        'Refactors code to be Pythonic (PEP 8 compliant)',
        'Optimizes for Apple Silicon M1 performance',
        'Improves readability and maintainability',
        'Removes technical debt and anti-patterns',
        'Applies best practices and design patterns'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('Pythonic Transformations:', style='Heading 3')
    pythonic = [
        'List comprehensions over loops',
        'Built-in functions (enumerate, zip, any, all)',
        'pathlib over os.path',
        'Type hints consistently',
        'Context managers (with statements)',
        'f-strings over .format()',
        'Dataclasses/Pydantic models'
    ]
    for item in pythonic:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('M1 Optimizations:', style='Heading 3')
    optimizations = [
        'Async/await for I/O operations',
        'Vectorized operations',
        'Minimize object creation',
        'Cache expensive computations',
        'Appropriate data structures',
        'Multiprocessing for CPU-bound tasks'
    ]
    for opt in optimizations:
        doc.add_paragraph(opt, style='List Bullet')
    
    doc.add_paragraph('Trigger: "Act as The Cleaner" or "Refactor this..."', style='Intense Quote')
    
    # 2.3 The Verifier
    add_heading_with_style(doc, '2.3 The Verifier (Agent B - Claude 3.5 Sonnet)', 2)
    doc.add_paragraph('Model: Claude 3.5 Sonnet (via Cursor)', style='Intense Quote')
    doc.add_paragraph('Focus: Testing and quality assurance', style='Intense Quote')
    
    doc.add_paragraph('Core Capabilities:', style='Heading 3')
    capabilities = [
        'Writes comprehensive test suites (pytest)',
        'Validates code correctness and edge cases',
        'Provides code quality feedback',
        'Checks integration points',
        'Verifies error handling'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('Test Coverage:', style='Heading 3')
    coverage = [
        'Unit tests (individual functions/methods)',
        'Integration tests (component interactions)',
        'Edge cases (empty inputs, None, boundaries)',
        'Error paths (exceptions, validation failures)',
        'Happy paths (normal operation flows)'
    ]
    for item in coverage:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Test Structure:', style='Heading 3')
    doc.add_paragraph(
        'Uses pytest with @pytest.mark.asyncio for async tests. '
        'Descriptive names: test_<function>_<scenario>_<expected_result>. '
        'Arrange-Act-Assert pattern. Independent, order-agnostic tests.'
    )
    
    doc.add_paragraph('Trigger: "Act as The Verifier" or "Test this..."', style='Intense Quote')
    
    # 2.4 The Detective
    add_heading_with_style(doc, '2.4 The Detective (Agent B - Claude 3.5 Sonnet)', 2)
    doc.add_paragraph('Model: Claude 3.5 Sonnet (via Cursor)', style='Intense Quote')
    doc.add_paragraph('Focus: Debugging and root cause analysis', style='Intense Quote')
    
    doc.add_paragraph('Core Capabilities:', style='Heading 3')
    capabilities = [
        'Investigates errors and exceptions systematically',
        'Traces bugs through the codebase',
        'Uses debugging tools effectively',
        'Identifies root causes (not just symptoms)',
        'Provides detailed bug reports with solutions'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('Debugging Methodology:', style='Heading 3')
    methodology = [
        '1. Gather information (errors, logs, recent changes)',
        '2. Reproduce the issue (minimal case)',
        '3. Trace execution (call stack, data flow)',
        '4. Use debugging tools (pdb, ipdb, logging, Docker logs)',
        '5. Root cause analysis (underlying issues)'
    ]
    for item in methodology:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Tools:', style='Heading 3')
    tools = [
        'Python: pdb, ipdb, logging, print statements',
        'Docker: docker compose logs -f [service]',
        'Database: SQL logging, connection checks',
        'Network: API endpoint verification'
    ]
    for tool in tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph('Trigger: "Act as The Detective" or "Debug this..."', style='Intense Quote')
    
    # 2.5 The EU Compliance Agent
    add_heading_with_style(doc, '2.5 The EU Compliance Agent (Agent A - Llama 3.1:8b)', 2)
    doc.add_paragraph('Model: Local Llama 3.1:8b via Ollama', style='Intense Quote')
    doc.add_paragraph('Focus: EU AI Act compliance auditing', style='Intense Quote')
    
    doc.add_paragraph('Core Capabilities:', style='Heading 3')
    capabilities = [
        'Audits code and AI system descriptions for EU AI Act compliance',
        'Classifies systems as High-Risk (Annex III) or not',
        'Verifies transparency requirements (Article 50)',
        'Checks for PII leakage and privacy violations',
        'Validates compliance with EU AI Act requirements',
        'Uses local Llama model for privacy-preserving audits'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('EU AI Act Knowledge:', style='Heading 3')
    doc.add_paragraph('High-Risk Systems (Annex III):', style='Heading 4')
    high_risk = [
        'Biometric identification and categorization',
        'Management and operation of critical infrastructure',
        'Educational and vocational training',
        'Employment, worker management, and access to self-employment',
        'Access to and enjoyment of essential private services and public services',
        'Law enforcement',
        'Migration, asylum, and border control management',
        'Administration of justice and democratic processes'
    ]
    for item in high_risk:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Key Requirements:', style='Heading 4')
    requirements = [
        'Article 50: AI-generated content must be detectable/flagged',
        'Article 10: Data governance (relevant, representative, error-free data)',
        'Article 13: User disclosure (inform users of AI interaction)',
        'Article 52: Deepfake labeling'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_paragraph('Audit Process:', style='Heading 3')
    audit_process = [
        '1. System classification (Annex III check)',
        '2. Compliance checks (transparency, disclosure, data governance)',
        '3. Privacy & security (PII leakage, data protection, access controls)'
    ]
    for item in audit_process:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Trigger: "Act as The EU Compliance Agent" or "Audit this for EU AI Act compliance..."', style='Intense Quote')
    
    # 2.6 The Coordinator
    add_heading_with_style(doc, '2.6 The Coordinator (Agent B - Claude 3.5 Sonnet)', 2)
    doc.add_paragraph('Model: Claude 3.5 Sonnet (via Cursor)', style='Intense Quote')
    doc.add_paragraph('Focus: Workflow orchestration and task planning', style='Intense Quote')
    
    doc.add_paragraph('Core Capabilities:', style='Heading 3')
    capabilities = [
        'Analyzes complex tasks and breaks them down into agent-specific steps',
        'Recommends which agents to use and in what sequence',
        'Orchestrates multi-agent workflows',
        'Tracks progress across agent interactions',
        'Provides workflow recommendations based on task requirements',
        'Optimizes agent selection for efficiency and quality'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('Trigger: "Act as The Coordinator" or "Plan a workflow for..."', style='Intense Quote')
    
    # 3. Model Comparison
    add_heading_with_style(doc, '3. Model Comparison', 1)
    
    model_data = [
        {'Agent': 'The Coder', 'Model': 'Llama 3.1:8b (local)', 'Speed': 'Fast', 'Best For': 'Quick code generation'},
        {'Agent': 'The Cleaner', 'Model': 'Claude 3.5 Sonnet', 'Speed': 'Medium', 'Best For': 'Complex refactoring'},
        {'Agent': 'The Verifier', 'Model': 'Claude 3.5 Sonnet', 'Speed': 'Medium', 'Best For': 'Thorough testing'},
        {'Agent': 'The Detective', 'Model': 'Claude 3.5 Sonnet', 'Speed': 'Medium', 'Best For': 'Deep debugging'},
        {'Agent': 'The EU Compliance Agent', 'Model': 'Llama 3.1:8b (local)', 'Speed': 'Fast', 'Best For': 'Privacy-preserving audits'},
        {'Agent': 'The Coordinator', 'Model': 'Claude 3.5 Sonnet', 'Speed': 'Medium', 'Best For': 'Planning & orchestration'}
    ]
    add_table_from_dict(doc, model_data)
    
    # 4. Recommended Workflows
    add_heading_with_style(doc, '4. Recommended Workflows', 1)
    
    doc.add_paragraph('Feature Development:', style='Heading 3')
    workflow1 = [
        '1. The Coder → Implement feature',
        '2. The Cleaner → Optimize and refactor',
        '3. The Verifier → Add tests',
        '4. The EU Compliance Agent → Audit compliance',
        '5. The Detective → Debug if needed'
    ]
    for step in workflow1:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph('Bug Fix:', style='Heading 3')
    workflow2 = [
        '1. The Detective → Investigate root cause',
        '2. The Coder → Implement fix',
        '3. The Cleaner → Ensure best practices',
        '4. The Verifier → Add regression tests'
    ]
    for step in workflow2:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph('Compliance Audit:', style='Heading 3')
    workflow3 = [
        '1. The EU Compliance Agent → Initial audit',
        '2. The Coder → Implement fixes if needed',
        '3. The Verifier → Validate fixes',
        '4. The EU Compliance Agent → Re-audit'
    ]
    for step in workflow3:
        doc.add_paragraph(step, style='List Bullet')
    
    # 5. Usage Guidelines
    add_heading_with_style(doc, '5. Usage Guidelines', 1)
    
    doc.add_paragraph('When to Use Each Agent:', style='Heading 3')
    usage = [
        'Need code quickly? → The Coder (Llama)',
        'Code needs improvement? → The Cleaner (Claude)',
        'Need tests? → The Verifier (Claude)',
        'Something broken? → The Detective (Claude)',
        'Compliance check? → The EU Compliance Agent (Llama)',
        'Complex multi-step task? → The Coordinator (me!) to plan the workflow'
    ]
    for item in usage:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph(
        'All agents work within your Docker environment and maintain context '
        'within conversations. Switch between agents by explicitly mentioning their role.'
    )
    
    # Footer
    doc.add_page_break()
    doc.add_paragraph('End of Document', style='Intense Quote')
    
    return doc


if __name__ == '__main__':
    print("Generating agent capabilities documentation...")
    doc = create_agent_documentation()
    output_path = '.cursor/Agent_Capabilities_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation saved to: {output_path}")
